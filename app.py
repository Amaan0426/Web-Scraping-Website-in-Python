from flask import Flask, render_template, request, redirect, url_for, flash, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField, SubmitField
from wtforms.validators import DataRequired, URL
from bs4 import BeautifulSoup
from textblob import TextBlob
from collections import Counter
import requests
import csv
from docx import Document
from openpyxl import Workbook
from scraper import scrape_data
from data_processing import process_data
import os
import logging
import psutil
from docx import Document

app = Flask(__name__, static_url_path='/static')
app.config['SECRET_KEY'] = 'ccea92330eaad6e31b676237a5d7a84e'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['STATIC_FOLDER'] = 'static'
db = SQLAlchemy(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(120), nullable=False)

class LoginForm(FlaskForm):
    username = StringField('Username', validators=[DataRequired()])
    password = PasswordField('Password', validators=[DataRequired()])
    submit = SubmitField('Login')

@login_manager.user_loader
def load_user(user_id):
    return db.session.get(User, int(user_id))


def get_top_keywords(text, num_keywords=5):
    words = text.split()
    return [word for word, _ in Counter(words).most_common(num_keywords)]

def arrange_data(data):
    arranged_data = []

    # Titles
    if 'titles' in data:
        section_header = "Titles:\n"
        titles_text = "\n".join(data['titles'])
        arranged_data.append(section_header + titles_text)

    # Paragraphs
    if 'paragraphs' in data:
        section_header = "\n\nParagraphs:\n"
        paragraphs_text = "\n".join(data['paragraphs'])
        arranged_data.append(section_header + paragraphs_text)

    # Links
    if 'links' in data:
        section_header = "\n\nLinks:\n"
        links_text = "\n".join(data['links'])
        arranged_data.append(section_header + links_text)

    # Lists
    if 'lists' in data:
        section_header = "\n\nLists:\n"
        lists_text = "\n".join(data['lists'])
        arranged_data.append(section_header + lists_text)

    # Tables
    if 'tables' in data:
        section_header = "\n\nTables:\n"
        # Handle tables based on your data structure

    # Images
    if 'images' in data:
        section_header = "\n\nImages:\n"
        images_text = "\n".join(data['images'])
        arranged_data.append(section_header + images_text)

    return "\n\n".join(arranged_data)

# Assume process_data just returns the scraped data for demonstration
def process_data(scraped_data):
    return scraped_data

def download_word(data):
    document = Document()
    document.add_heading('Scraped Data', level=1)

    # Arrange the data into sections
    arranged_data = arrange_data(data)

    # Add each section to the document
    for section_title, section_content in arranged_data.items():
        if section_content:
            document.add_heading(section_title, level=2)
            document.add_paragraph(section_content)

    filename = 'scraped_data.docx'
    document.save(filename)
    return filename

def download_excel(data):
    workbook = Workbook()
    sheet = workbook.active
    
    # Arrange the data into sections
    arranged_data = arrange_data(data)

    # Add each section to the Excel sheet
    for i, (section_title, section_content) in enumerate(arranged_data.items(), start=1):
        if section_content:
            sheet.cell(row=i, column=1, value=section_title)
            sheet.cell(row=i, column=2, value=section_content)

    filename = 'scraped_data.xlsx'
    workbook.save(filename)
    return filename

def download_csv(data):
    filename = 'scraped_data.csv'
    
    # Arrange the data into sections
    arranged_data = arrange_data(data)

    # Write the arranged data to a CSV file
    with open(filename, 'w', newline='', encoding='utf-8') as csv_file:
        csv_writer = csv.writer(csv_file)
        for section_title, section_content in arranged_data.items():
            if section_content:
                csv_writer.writerow([section_title, section_content])

    return filename

def is_file_open(file_path):
    try:
        process = psutil.Process(os.getpid())
        open_files = process.open_files()
        
        if open_files is not None and file_path in [file.path for file in open_files]:
            return True
    except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.ZombieProcess):
        # Handle exceptions raised by psutil (e.g., if the process does not exist)
        pass

    return False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        if User.query.filter_by(username=username).first():
            flash('Username already exists. Please choose a different one.', 'error')
            return redirect(url_for('register'))

        new_user = User(username=username, password=password)
        db.session.add(new_user)
        db.session.commit()

        flash('Account created successfully. Please log in.', 'success')
        return redirect(url_for('login'))

    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()

    if form.validate_on_submit():
        username = form.username.data
        password = form.password.data

        user = User.query.filter_by(username=username).first()

        if user and user.password == password:
            login_user(user)
            flash('Logged in successfully.', 'success')
            return redirect(url_for('dashboard'))  # Redirect to the dashboard

        flash('Invalid username or password. Please try again.', 'error')

    return render_template('login.html', form=form)

@app.route('/dashboard')
@login_required
def dashboard():
    return render_template('dashboard.html', scraped_data=None)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('Logged out successfully.', 'success')
    return redirect(url_for('index'))

# Sample data for demonstration
sample_permissions = [
    {'website': 'example.com', 'status': 'Pending'},
    {'website': 'sample.com', 'status': 'Granted'},
]


@app.route('/scrape', methods=['POST'])
def scrape():
    try:
        if request.is_json:
            data = request.get_json()
            url = data.get('url', '')
        else:
            url = request.form.get('url', '')

        # Step 1: Scrape the data
        scraped_data = scrape_data(url)

        if isinstance(scraped_data, str):
            # If there's an error during scraping, log and return an error response
            logging.error(f'Scraping error: {scraped_data}')
            flash(scraped_data, 'error')
            return jsonify({'error': scraped_data})

        # Step 2: Process the scraped data if needed
        processed_data = process_data(scraped_data)

        if processed_data is None:
            # If there's an issue during processing, log and return an error response
            error_message = 'Error during data processing.'
            logging.error(error_message)
            flash(error_message, 'error')
            return jsonify({'error': error_message})

        # Log the processed data (optional)
        logging.info(f'Processed Data: {processed_data}')

        # Check if processed_data is a dictionary, and convert it to a string if needed
        if not isinstance(processed_data, str):
            processed_data = str(processed_data)

        # Create a temporary file to store the processed data
        temp_filename = 'temp_data.txt'
        with open(temp_filename, 'w', encoding='utf-8') as temp_file:
            temp_file.write(processed_data)

        # Return a downloadable file response
        return send_file(temp_filename, as_attachment=True, download_name='scraped_data.txt')

    except Exception as e:
        # Handle exceptions and return an error response
        error_message = f'Error: {str(e)}'
        logging.error(error_message)
        return jsonify({'data': None, 'error': error_message}), 500

    finally:
        # Cleanup: Remove the temporary file if it exists and is closed
        if os.path.exists(temp_filename) and not is_file_open(temp_filename):
            os.remove(temp_filename)



@app.route('/download', methods=['POST'])
@login_required
def download():
    url = request.form['url']
    try:
        response = requests.get(url, timeout=5)
        response.raise_for_status()
    except requests.exceptions.RequestException as err:
        return jsonify({'error': f"An error occurred: {err}"})

    soup = BeautifulSoup(response.text, 'html.parser')
    scraped_data = soup.get_text()

    format_type = request.form['format']

    if format_type == 'word':
        download_word(scraped_data)
        return send_file('scraped_data.docx', as_attachment=True)
    elif format_type == 'excel':
        download_excel(scraped_data)
        return send_file('scraped_data.xlsx', as_attachment=True)
    elif format_type == 'csv':
        download_csv(scraped_data)
        return send_file('scraped_data.csv', as_attachment=True)
    else:
        return jsonify({'error': 'Invalid format type'})

@app.route('/permissions')
@login_required
def permissions():
    return render_template('permissions.html', permissions=sample_permissions)

if __name__ == '__main__':
    app.run(debug=True)
